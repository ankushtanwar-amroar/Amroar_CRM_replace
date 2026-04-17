"""
CLU-BOT External Context MCP Service
Handles file content extraction, URL fetching, and context-aware analysis.
All external data is safely parsed and truncated before LLM consumption.
"""

import os
import re
import csv
import io
import ipaddress
import logging
from datetime import datetime, timezone
from typing import Dict, Any, Optional, List
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup

from ..models import (
    ReadFilePayload, FetchUrlPayload, AnalyzeWithContextPayload,
    ALLOWED_FILE_TYPES, MAX_CONTENT_CHARS
)
import google.generativeai as genai

logger = logging.getLogger(__name__)

# SSRF protection: blocked hosts and IP ranges
BLOCKED_HOSTS = {"localhost", "127.0.0.1", "0.0.0.0", "::1", "metadata.google.internal"}
BLOCKED_IP_RANGES = [
    ipaddress.ip_network("10.0.0.0/8"),
    ipaddress.ip_network("172.16.0.0/12"),
    ipaddress.ip_network("192.168.0.0/16"),
    ipaddress.ip_network("169.254.0.0/16"),
    ipaddress.ip_network("127.0.0.0/8"),
    ipaddress.ip_network("::1/128"),
    ipaddress.ip_network("fc00::/7"),
    ipaddress.ip_network("fe80::/10"),
]


class ExternalContextMCPService:
    """
    MCP Service for external context operations.
    - Read uploaded files (PDF, CSV, DOCX, TXT, XLSX)
    - Fetch and parse URL content
    - Combine CRM data with external context for LLM analysis
    """

    def __init__(self, db):
        self.db = db

    @staticmethod
    def _extract_query_terms(query: str) -> List[str]:
        """Extract meaningful search terms from a user query."""
        if not query:
            return []
        terms: List[str] = []
        for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9\-]{1,}", query):
            t = token.strip()
            if not t:
                continue
            # Keep acronyms like IREDA and normal words with useful length.
            if t.isupper() and len(t) >= 2:
                terms.append(t.lower())
            elif len(t) >= 3:
                terms.append(t.lower())
        # Preserve order while de-duplicating
        seen = set()
        deduped: List[str] = []
        for t in terms:
            if t not in seen:
                seen.add(t)
                deduped.append(t)
        return deduped

    def _build_query_focused_context(self, content: str, query: str, max_chars: int) -> tuple[str, bool]:
        """
        Build a query-focused context window from long content.
        Returns (selected_context, truncated_flag).
        """
        if not content:
            return "", False
        if len(content) <= max_chars:
            return content, False

        terms = self._extract_query_terms(query)
        chunk_size = 2200
        overlap = 250
        step = max(500, chunk_size - overlap)
        chunks: List[tuple[int, str]] = []
        idx = 0
        for start in range(0, len(content), step):
            chunk = content[start:start + chunk_size]
            if not chunk:
                break
            chunks.append((idx, chunk))
            idx += 1
            if start + chunk_size >= len(content):
                break

        def score_chunk(text: str) -> int:
            if not terms:
                return 0
            lower_text = text.lower()
            score = 0
            for term in terms:
                # Count exact term hits and give slight bonus to early hits.
                hits = lower_text.count(term)
                if hits:
                    score += hits * 10
                    first_pos = lower_text.find(term)
                    if 0 <= first_pos < 300:
                        score += 2
            return score

        scored: List[tuple[int, int]] = []
        for chunk_idx, chunk_text in chunks:
            scored.append((chunk_idx, score_chunk(chunk_text)))

        scored.sort(key=lambda x: (-x[1], x[0]))
        selected_indexes: List[int] = []
        used = 0

        # Prefer chunks that actually match the query terms.
        for chunk_idx, sc in scored:
            if sc <= 0:
                continue
            chunk_len = len(chunks[chunk_idx][1])
            if used + chunk_len > max_chars and selected_indexes:
                continue
            selected_indexes.append(chunk_idx)
            used += chunk_len
            if used >= max_chars:
                break

        # Fallback: include beginning + end snippets if no lexical match.
        if not selected_indexes:
            first_chunk = chunks[0][1][: max_chars // 2]
            last_chunk = chunks[-1][1][: max_chars - len(first_chunk)]
            merged = f"{first_chunk}\n\n... (content omitted) ...\n\n{last_chunk}"
            return merged[:max_chars], True

        selected_indexes.sort()
        selected_text_parts: List[str] = []
        remaining = max_chars
        for chunk_idx in selected_indexes:
            chunk_text = chunks[chunk_idx][1]
            if len(chunk_text) > remaining:
                chunk_text = chunk_text[:remaining]
            selected_text_parts.append(chunk_text)
            remaining -= len(chunk_text)
            if remaining <= 0:
                break

        merged = "\n\n... (content omitted) ...\n\n".join(selected_text_parts)
        return merged[:max_chars], True

    # =========================================================================
    # File Content Extraction
    # =========================================================================

    async def read_file(
        self,
        tenant_id: str,
        user_id: str,
        payload: ReadFilePayload
    ) -> Dict[str, Any]:
        """Read and extract content from an uploaded file"""
        try:
            file_record = await self.db.clu_bot_file_uploads.find_one(
                {"id": payload.file_id, "tenant_id": tenant_id},
                {"_id": 0}
            )

            if not file_record:
                return {
                    "success": False,
                    "message": "File not found. Please upload a file first using the attach button."
                }

            storage_path = file_record.get("storage_path")
            file_type = file_record.get("file_type", "").lower()
            file_name = file_record.get("file_name", "unknown")

            if not storage_path or not os.path.exists(storage_path):
                return {
                    "success": False,
                    "message": f"File '{file_name}' storage is unavailable. Please re-upload."
                }

            content = self._extract_file_content(storage_path, file_type)

            if not content or not content.strip():
                return {
                    "success": False,
                    "message": f"Could not extract text content from '{file_name}'. The file may be empty or in an unsupported format."
                }

            full_length = len(content)
            truncated = full_length > MAX_CONTENT_CHARS

            # If user has a specific query, use LLM to answer it
            if payload.query:
                focused_content, focused_truncated = self._build_query_focused_context(
                    content=content,
                    query=payload.query,
                    max_chars=MAX_CONTENT_CHARS
                )
                answer = await self._llm_analyze_content(
                    content=focused_content,
                    query=payload.query,
                    source_label=f"file '{file_name}'"
                )
                return {
                    "success": True,
                    "file_name": file_name,
                    "file_type": file_type,
                    "content_length": full_length,
                    "truncated": focused_truncated,
                    "query": payload.query,
                    "answer": answer,
                    "summary": answer
                }

            # Return a summary of the content
            content = content[:MAX_CONTENT_CHARS]
            summary = await self._llm_summarize_content(
                content=content,
                source_label=f"file '{file_name}' ({file_type.upper()})"
            )

            return {
                "success": True,
                "file_name": file_name,
                "file_type": file_type,
                "content_length": full_length,
                "truncated": truncated,
                "summary": summary
            }

        except Exception as e:
            logger.error(f"Read file error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to read file: {str(e)}"
            }

    def _extract_file_content(self, file_path: str, file_type: str) -> str:
        """Extract text content from a file based on its type"""

        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        elif file_type == "csv":
            return self._extract_csv(file_path)
        elif file_type == "xlsx":
            return self._extract_xlsx(file_path)
        elif file_type == "txt":
            return self._extract_txt(file_path)
        else:
            return ""

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))

            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""

    def _extract_csv(self, file_path: str) -> str:
        """Extract content from CSV as structured text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                rows = list(reader)

            if not rows:
                return ""

            # Format as a readable table
            lines = []
            headers = rows[0] if rows else []
            lines.append(" | ".join(headers))
            lines.append("-" * len(lines[0]))

            for row in rows[1:100]:  # Limit to 100 rows
                lines.append(" | ".join(row))

            if len(rows) > 101:
                lines.append(f"... ({len(rows) - 101} more rows)")

            return "\n".join(lines)
        except Exception as e:
            logger.error(f"CSV extraction error: {e}")
            return ""

    def _extract_xlsx(self, file_path: str) -> str:
        """Extract content from XLSX"""
        try:
            from openpyxl import load_workbook
            wb = load_workbook(file_path, read_only=True, data_only=True)
            all_text = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_text.append(f"--- Sheet: {sheet_name} ---")

                row_count = 0
                for row in ws.iter_rows(values_only=True):
                    if row_count >= 100:
                        all_text.append("... (more rows in sheet)")
                        break
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        all_text.append(" | ".join(cells))
                        row_count += 1

            wb.close()
            return "\n".join(all_text)
        except Exception as e:
            logger.error(f"XLSX extraction error: {e}")
            return ""

    def _extract_txt(self, file_path: str) -> str:
        """Extract plain text"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""

    # =========================================================================
    # URL Content Fetching
    # =========================================================================

    async def fetch_url(
        self,
        tenant_id: str,
        user_id: str,
        payload: FetchUrlPayload
    ) -> Dict[str, Any]:
        """Fetch and parse content from a URL with SSRF protection"""
        try:
            url = payload.url.strip()

            # Validate URL
            validation_error = self._validate_url(url)
            if validation_error:
                return {"success": False, "message": validation_error}

            # Fetch content
            async with httpx.AsyncClient(
                timeout=15.0,
                follow_redirects=True,
                max_redirects=3,
                verify=False
            ) as client:
                response = await client.get(url, headers={
                    "User-Agent": "CLU-BOT/1.0 (CRM Assistant)"
                })

            if response.status_code != 200:
                return {
                    "success": False,
                    "message": f"Failed to fetch URL (HTTP {response.status_code}). The site may be unavailable or blocking automated access."
                }

            content_type = response.headers.get("content-type", "")

            # Parse HTML content
            if "html" in content_type or "text" in content_type:
                text = self._parse_html_content(response.text)
            else:
                text = response.text[:MAX_CONTENT_CHARS]

            if not text or not text.strip():
                return {
                    "success": False,
                    "message": "Could not extract readable content from this URL."
                }

            truncated = len(text) > MAX_CONTENT_CHARS
            text = text[:MAX_CONTENT_CHARS]

            # If user has a specific query, use LLM
            if payload.query:
                answer = await self._llm_analyze_content(
                    content=text,
                    query=payload.query,
                    source_label=f"URL '{url}'"
                )
                return {
                    "success": True,
                    "url": url,
                    "content_length": len(text),
                    "truncated": truncated,
                    "query": payload.query,
                    "answer": answer,
                    "summary": answer
                }

            summary = await self._llm_summarize_content(
                content=text,
                source_label=f"URL '{url}'"
            )

            return {
                "success": True,
                "url": url,
                "content_length": len(text),
                "truncated": truncated,
                "summary": summary
            }

        except httpx.TimeoutException:
            return {
                "success": False,
                "message": "URL request timed out. The site may be slow or unresponsive."
            }
        except Exception as e:
            logger.error(f"Fetch URL error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to fetch URL: {str(e)}"
            }

    def _validate_url(self, url: str) -> Optional[str]:
        """Validate URL for safety (SSRF prevention)"""
        try:
            parsed = urlparse(url)
        except Exception:
            return "Invalid URL format."

        # Only allow HTTP/HTTPS
        if parsed.scheme not in ("http", "https"):
            return f"Only HTTP and HTTPS URLs are allowed. Got: {parsed.scheme}://"

        if not parsed.hostname:
            return "URL must have a valid hostname."

        hostname = parsed.hostname.lower()

        # Block known dangerous hosts
        if hostname in BLOCKED_HOSTS:
            return "This URL points to a restricted address and cannot be accessed."

        # Block internal IP ranges
        try:
            ip = ipaddress.ip_address(hostname)
            for network in BLOCKED_IP_RANGES:
                if ip in network:
                    return "This URL points to an internal network address and cannot be accessed."
        except ValueError:
            pass  # hostname is a domain name, not IP — that's fine

        return None

    def _parse_html_content(self, html: str) -> str:
        """Parse HTML and extract clean text content. Strip scripts/styles."""
        soup = BeautifulSoup(html, 'html.parser')

        # Remove script and style elements
        for element in soup(['script', 'style', 'noscript', 'iframe', 'svg',
                            'nav', 'footer', 'header']):
            element.decompose()

        # Extract text
        text = soup.get_text(separator='\n', strip=True)

        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        text = '\n'.join(lines)

        # Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text

    # =========================================================================
    # Context-Aware Analysis (CRM + External)
    # =========================================================================

    async def analyze_with_context(
        self,
        tenant_id: str,
        user_id: str,
        payload: AnalyzeWithContextPayload
    ) -> Dict[str, Any]:
        """Combine CRM data with external context for LLM analysis"""
        try:
            context_parts = []
            sources_used = []

            # 1. Fetch external file content if file_id provided
            if payload.file_id:
                file_record = await self.db.clu_bot_file_uploads.find_one(
                    {"id": payload.file_id, "tenant_id": tenant_id},
                    {"_id": 0}
                )
                if file_record and os.path.exists(file_record.get("storage_path", "")):
                    file_content = self._extract_file_content(
                        file_record["storage_path"],
                        file_record.get("file_type", "txt")
                    )
                    if file_content:
                        context_parts.append(f"=== FILE: {file_record.get('file_name', 'uploaded file')} ===\n{file_content[:MAX_CONTENT_CHARS // 2]}")
                        sources_used.append(f"File: {file_record.get('file_name')}")

            # 2. Fetch URL content if url provided
            if payload.url:
                validation_error = self._validate_url(payload.url)
                if not validation_error:
                    try:
                        async with httpx.AsyncClient(timeout=10.0, follow_redirects=True, max_redirects=3, verify=False) as client:
                            resp = await client.get(payload.url, headers={"User-Agent": "CLU-BOT/1.0"})
                        if resp.status_code == 200:
                            ct = resp.headers.get("content-type", "")
                            url_text = self._parse_html_content(resp.text) if "html" in ct else resp.text
                            context_parts.append(f"=== URL: {payload.url} ===\n{url_text[:MAX_CONTENT_CHARS // 2]}")
                            sources_used.append(f"URL: {payload.url}")
                    except Exception as e:
                        logger.warning(f"URL fetch for context failed: {e}")

            # 3. Fetch CRM data if object_type provided
            if payload.crm_object_type:
                query_filter = {"tenant_id": tenant_id, "object_name": payload.crm_object_type, "is_deleted": {"$ne": True}}
                if payload.crm_search_term:
                    query_filter["$or"] = [
                        {"data.first_name": {"$regex": payload.crm_search_term, "$options": "i"}},
                        {"data.last_name": {"$regex": payload.crm_search_term, "$options": "i"}},
                        {"data.name": {"$regex": payload.crm_search_term, "$options": "i"}},
                        {"data.company": {"$regex": payload.crm_search_term, "$options": "i"}},
                        {"data.opportunity_name": {"$regex": payload.crm_search_term, "$options": "i"}},
                        {"data.account_name": {"$regex": payload.crm_search_term, "$options": "i"}},
                    ]

                cursor = self.db.object_records.find(query_filter, {"_id": 0}).limit(10)
                records = await cursor.to_list(10)

                if records:
                    crm_lines = [f"=== CRM DATA: {payload.crm_object_type.title()} Records ==="]
                    for r in records:
                        data = r.get("data", {})
                        name = data.get("name") or f"{data.get('first_name', '')} {data.get('last_name', '')}".strip() or data.get("opportunity_name", "Unknown")
                        fields = ", ".join(f"{k}: {v}" for k, v in data.items() if v and k not in ("id",))
                        crm_lines.append(f"- {name}: {fields[:300]}")
                    context_parts.append("\n".join(crm_lines))
                    sources_used.append(f"CRM: {len(records)} {payload.crm_object_type} records")

            if not context_parts:
                return {
                    "success": False,
                    "message": "No external context could be gathered. Please upload a file, provide a URL, or specify CRM data to analyze."
                }

            # Combine and send to LLM
            combined_context = "\n\n".join(context_parts)
            # Truncate combined context
            combined_context = combined_context[:MAX_CONTENT_CHARS]

            analysis = await self._llm_analyze_content(
                content=combined_context,
                query=payload.query,
                source_label=f"combined context ({', '.join(sources_used)})"
            )

            return {
                "success": True,
                "sources": sources_used,
                "query": payload.query,
                "analysis": analysis,
                "summary": analysis
            }

        except Exception as e:
            logger.error(f"Analyze with context error: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "message": f"Failed to analyze with context: {str(e)}"
            }

    # =========================================================================
    # LLM Helper Methods
    # =========================================================================

    # async def _llm_summarize_content(self, content: str, source_label: str) -> str:
    #     """Use LLM to summarize extracted content"""
    #     try:
    #         from emergentintegrations.llm.chat import LlmChat, UserMessage
    #         import uuid

    #         chat = LlmChat(
    #             api_key=os.environ.get("EMERGENT_LLM_KEY", ""),
    #             session_id=f"clubot-ctx-{uuid.uuid4().hex[:8]}",
    #             system_message=(
    #                 "You are a CRM assistant that summarizes documents and web content concisely. "
    #                 "Provide a structured summary with key points. Be specific and actionable. "
    #                 "Format using bullet points and bold headers where helpful."
    #             )
    #         ).with_model("gemini", "gemini-2.5-flash")

    #         response = await chat.send_message(
    #             UserMessage(text=f"Summarize the following content from {source_label}:\n\n{content}")
    #         )
    #         return response

    #     except Exception as e:
    #         logger.error(f"LLM summarize error: {e}")
    #         return f"**Content from {source_label}** (auto-extract, LLM unavailable):\n\n{content[:2000]}"



    # async def _llm_analyze_content(self, content: str, query: str, source_label: str) -> str:
    #     """Use LLM to analyze content against a specific query"""
    #     try:
    #         from emergentintegrations.llm.chat import LlmChat, UserMessage
    #         import uuid

    #         chat = LlmChat(
    #             api_key=os.environ.get("EMERGENT_LLM_KEY", ""),
    #             session_id=f"clubot-ctx-{uuid.uuid4().hex[:8]}",
    #             system_message=(
    #                 "You are a CRM assistant that analyzes documents, web content, and CRM data. "
    #                 "Answer the user's question based on the provided context. Be specific, "
    #                 "accurate, and actionable. If the context doesn't contain enough information "
    #                 "to fully answer, say so clearly."
    #             )
    #         ).with_model("gemini", "gemini-2.5-flash")

    #         response = await chat.send_message(
    #             UserMessage(text=f"Context from {source_label}:\n\n{content}\n\n---\n\nQuestion: {query}")
    #         )
    #         return response

    #     except Exception as e:
    #         logger.error(f"LLM analyze error: {e}")
    #         return f"Unable to analyze content from {source_label} at this time. Error: {str(e)}"

    async def _llm_summarize_content(self, content: str, source_label: str) -> str:

        """Use Gemini LLM to summarize extracted content"""
        try:

            # Configure API
            genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))

            model = genai.GenerativeModel("gemini-2.5-flash")

            prompt = f"""
                You are a CRM assistant that summarizes documents and web content concisely.

                Provide:
                - Structured summary
                - Key points
                - Actionable insights

                Use bullet points and bold headers.

                Content source: {source_label}

                Content:
                {content}
                """

            response = await model.generate_content_async(prompt)

            return response.text

        except Exception as e:
            logger.error(f"Gemini summarize error: {e}")
            return f"**Content from {source_label}** (auto-extract, LLM unavailable):\n\n{content[:2000]}"
    
    async def _llm_analyze_content(self, content: str, query: str, source_label: str) -> str:

        """Use Gemini LLM to analyze content against a specific query"""
        try:
            # Configure API
            genai.configure(api_key=os.environ.get("GEMINI_API_KEY"))

            model = genai.GenerativeModel("gemini-2.5-flash")

            prompt = f"""
            You are a CRM assistant that analyzes documents, web content, and CRM data.

            Instructions:
            - Answer the user's question using ONLY the provided context
            - Be specific, accurate, and actionable
            - If the answer is not fully available in the context, clearly say what is missing
            - Prefer structured output (bullet points if helpful)

            Context source: {source_label}

            Context:
            {content}

            ---

            User Question:
            {query}
            """

            generation_config = {
                "temperature": 0.3,
                "top_p": 0.9,
                "max_output_tokens": 1024,
            }

            response = await model.generate_content_async(
                prompt,
                generation_config=generation_config
            )

            return response.text

        except Exception as e:
            logger.error(f"Gemini analyze error: {e}")
            return f"Unable to analyze content from {source_label} at this time. Error: {str(e)}"
    # =========================================================================
    # File Upload Management
    # =========================================================================

    async def get_user_files(self, tenant_id: str, user_id: str) -> List[Dict[str, Any]]:
        """Get list of files uploaded by this user for CLU-BOT context"""
        uid = str(user_id) if user_id is not None else user_id
        cursor = self.db.clu_bot_file_uploads.find(
            {"tenant_id": tenant_id, "user_id": uid},
            {"_id": 0}
        ).sort("uploaded_at", -1).limit(20)
        return await cursor.to_list(20)

    async def get_latest_file(self, tenant_id: str, user_id: str) -> Optional[Dict[str, Any]]:
        """Get the most recently uploaded file for this user"""
        uid = str(user_id) if user_id is not None else user_id
        doc = await self.db.clu_bot_file_uploads.find_one(
            {"tenant_id": tenant_id, "user_id": uid},
            {"_id": 0},
            sort=[("uploaded_at", -1)]
        )
        if doc or not user_id:
            return doc
        # Legacy rows may have stored user_id as non-string (e.g. UUID type)
        return await self.db.clu_bot_file_uploads.find_one(
            {"tenant_id": tenant_id, "user_id": user_id},
            {"_id": 0},
            sort=[("uploaded_at", -1)]
        )


# Singleton
_external_context_service = None


def get_external_context_mcp_service(db) -> ExternalContextMCPService:
    """Get ExternalContextMCPService instance"""
    global _external_context_service
    if _external_context_service is None:
        _external_context_service = ExternalContextMCPService(db)
    return _external_context_service
