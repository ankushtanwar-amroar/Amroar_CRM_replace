"""
Document Conversion Service — Converts PDF/DOC/DOCX to structured editable content blocks.
All uploaded documents are converted to the SAME block format as AI-generated templates,
ensuring fully editable, persistent, and consistent editing experience.
"""
import logging
import io
import re
import uuid
import hashlib
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)


class DocumentConversionService:
    """Converts uploaded documents to structured editable content blocks."""

    def convert(self, file_bytes: bytes, filename: str, s3_service=None, tenant_id=None, template_id=None) -> Dict[str, Any]:
        """Convert a document file to structured content blocks."""
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

        if ext in ('docx',):
            blocks = self.convert_docx_to_blocks(file_bytes, s3_service, tenant_id, template_id)
            source_format = "docx"
        elif ext in ('doc',):
            blocks = [{
                "id": "block_0",
                "type": "paragraph",
                "content": "This .doc file needs to be saved as .docx for full editing support.",
                "page": 1,
                "style": {"fontSize": "12px"},
                "editable": True,
            }]
            source_format = "doc"
        elif ext == 'pdf':
            blocks = self.convert_pdf_to_blocks(file_bytes, s3_service, tenant_id, template_id)
            source_format = "pdf"
        else:
            raise ValueError(f"Unsupported file type: .{ext}")

        return {
            "blocks": blocks,
            "source_format": source_format,
            "block_count": len(blocks),
        }

    # ── PDF Conversion ────────────────────────────────────────────────

    def convert_pdf_to_blocks(self, file_bytes: bytes, s3_service=None, tenant_id=None, template_id=None) -> List[Dict[str, Any]]:
        """Convert PDF to structured content blocks (same format as AI templates).
        
        Uses pdfminer layout analysis for font-aware extraction, producing
        proper heading/paragraph blocks with formatting metadata.
        Extracts embedded images and uploads to S3.
        """
        # Primary: pdfminer with font analysis for rich block extraction
        blocks = self._extract_rich_blocks_pdfminer(file_bytes)

        # Fallback: PyPDF2 basic extraction
        if not blocks:
            logger.info("pdfminer produced no blocks, falling back to PyPDF2")
            blocks = self._extract_blocks_pypdf2(file_bytes)

        # Extract and insert images from the PDF
        if s3_service and tenant_id:
            image_blocks = self._extract_pdf_images(file_bytes, s3_service, tenant_id, template_id)
            if image_blocks:
                blocks = self._merge_image_blocks(blocks, image_blocks)

        return blocks

    def _extract_rich_blocks_pdfminer(self, file_bytes: bytes) -> List[Dict[str, Any]]:
        """Extract structured blocks from PDF using pdfminer's layout analysis.
        
        Analyzes font sizes to determine headings vs paragraphs,
        detects bold/italic from font names, and preserves alignment.
        """
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTTextBox, LTTextLine, LTChar, LAParams

        laparams = LAParams(
            line_margin=0.5,
            word_margin=0.1,
            char_margin=2.0,
            boxes_flow=0.5,
        )

        try:
            page_layouts = list(extract_pages(io.BytesIO(file_bytes), laparams=laparams))
        except Exception as e:
            logger.warning(f"pdfminer extraction failed: {e}")
            return []

        # First pass: collect all font sizes to determine what's a heading
        all_font_sizes = []
        for page_layout in page_layouts:
            for element in page_layout:
                if isinstance(element, LTTextBox):
                    for line in element:
                        if isinstance(line, LTTextLine):
                            for char in line:
                                if isinstance(char, LTChar) and char.size > 0:
                                    all_font_sizes.append(round(char.size, 1))

        if not all_font_sizes:
            return []

        # Determine font size thresholds
        from collections import Counter
        size_counts = Counter(all_font_sizes)
        most_common_size = size_counts.most_common(1)[0][0] if size_counts else 12
        heading_threshold = most_common_size * 1.3  # 30% larger = heading
        subheading_threshold = most_common_size * 1.1  # 10% larger = subheading

        blocks = []
        for page_num, page_layout in enumerate(page_layouts, 1):
            page_w = page_layout.width or 612

            for element in page_layout:
                if not isinstance(element, LTTextBox):
                    continue

                text = element.get_text().strip()
                if not text:
                    continue

                # Analyze font characteristics from the text box
                font_size = most_common_size
                is_bold = False
                is_italic = False
                char_count = 0

                for line in element:
                    if isinstance(line, LTTextLine):
                        for char in line:
                            if isinstance(char, LTChar):
                                char_count += 1
                                if char_count <= 5:  # Sample first few chars
                                    font_size = max(font_size, round(char.size, 1))
                                fname = (char.fontname or "").lower()
                                if "bold" in fname or "black" in fname or "heavy" in fname:
                                    is_bold = True
                                if "italic" in fname or "oblique" in fname:
                                    is_italic = True

                # Determine block type from font size
                if font_size >= heading_threshold:
                    block_type = "heading"
                elif font_size >= subheading_threshold and is_bold:
                    block_type = "subheading"
                elif is_bold and len(text) < 80 and '\n' not in text.strip():
                    block_type = "subheading"
                else:
                    block_type = "paragraph"

                # Detect alignment from position
                x0 = element.bbox[0]
                x1 = element.bbox[2]
                center_x = (x0 + x1) / 2
                text_align = "left"
                if abs(center_x - page_w / 2) < page_w * 0.08:
                    text_align = "center"
                elif x0 > page_w * 0.6:
                    text_align = "right"

                # Clean up text content
                content = text.replace('\n', ' ').strip()
                # Preserve intentional line breaks for form fields
                if '___' in text or '____' in text:
                    content = text.strip()

                blocks.append({
                    "id": f"blk_{page_num}_{len(blocks)}",
                    "type": block_type,
                    "content": content,
                    "page": page_num,
                    "style": {
                        "fontSize": f"{int(font_size)}px",
                        "fontWeight": "bold" if is_bold or block_type in ("heading", "subheading") else "normal",
                        "fontStyle": "italic" if is_italic else "normal",
                        "textAlign": text_align,
                    },
                    "editable": True,
                })

        logger.info(f"pdfminer extracted {len(blocks)} rich blocks from {len(page_layouts)} pages")
        return blocks

    def _extract_blocks_pypdf2(self, file_bytes: bytes) -> List[Dict[str, Any]]:
        """Fallback extraction using PyPDF2 (no font info, basic paragraph splitting)."""
        from PyPDF2 import PdfReader

        blocks = []
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
        except Exception as e:
            logger.warning(f"PyPDF2 failed to read PDF: {e}")
            return []

        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if not text.strip():
                continue

            paragraphs = self._split_into_paragraphs(text)
            for p_idx, para in enumerate(paragraphs):
                para = para.strip()
                if not para:
                    continue

                lines = para.split('\n')
                is_heading = (
                    len(lines) == 1
                    and len(para) < 80
                    and (para.isupper() or para.istitle())
                )

                block_type = "heading" if is_heading and p_idx < 3 else (
                    "subheading" if is_heading else "paragraph"
                )

                blocks.append({
                    "id": f"blk_{page_num}_{p_idx}",
                    "type": block_type,
                    "content": para.replace('\n', ' ').strip(),
                    "page": page_num,
                    "style": {
                        "fontSize": "20px" if block_type == "heading"
                                    else "16px" if block_type == "subheading"
                                    else "12px",
                        "fontWeight": "bold" if block_type in ("heading", "subheading") else "normal",
                        "textAlign": "left",
                    },
                    "editable": True,
                })

        return blocks

    @staticmethod
    def _split_into_paragraphs(text: str) -> List[str]:
        """Smart paragraph splitting that handles PDFs without double-newlines."""
        parts = re.split(r'\n\s*\n', text)
        if len(parts) > 1:
            return [p.strip() for p in parts if p.strip()]

        lines = text.split('\n')
        paragraphs = []
        current = []

        section_start = re.compile(
            r'^(?:'
            r'\d+\.\s'
            r'|[A-Z][A-Z\s]{4,}$'
            r'|\([a-z]\)\s'
            r'|(?:Party\s[AB]:)'
            r'|(?:WHEREAS|NOW,?\s+THEREFORE|IN WITNESS)'
            r')'
        )

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if current:
                    paragraphs.append('\n'.join(current))
                    current = []
                continue
            if section_start.match(stripped) and current:
                paragraphs.append('\n'.join(current))
                current = [stripped]
            else:
                current.append(stripped)

        if current:
            paragraphs.append('\n'.join(current))
        return [p.strip() for p in paragraphs if p.strip()]

    # ── PDF Image Extraction ──────────────────────────────────────────

    def _extract_pdf_images(self, file_bytes: bytes, s3_service, tenant_id: str, template_id: str = None) -> List[Dict[str, Any]]:
        """Extract images from PDF pages and upload to S3 as image blocks.
        Uses PyMuPDF (fitz) as primary, falls back to PyPDF2.
        """
        # Try PyMuPDF first (more reliable)
        image_blocks = self._extract_images_pymupdf(file_bytes, s3_service, tenant_id, template_id)
        if image_blocks:
            return image_blocks

        # Fallback to PyPDF2
        return self._extract_images_pypdf2(file_bytes, s3_service, tenant_id, template_id)

    def _extract_images_pymupdf(self, file_bytes: bytes, s3_service, tenant_id: str, template_id: str = None) -> List[Dict[str, Any]]:
        """Extract images using PyMuPDF (fitz) — handles most PDF image encodings."""
        try:
            import fitz
        except ImportError:
            logger.info("PyMuPDF not installed, skipping fitz-based image extraction")
            return []

        image_blocks = []
        seen_hashes = set()

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as e:
            logger.warning(f"PyMuPDF failed to open PDF: {e}")
            return []

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_idx, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if not base_image:
                        continue

                    img_data = base_image.get("image")
                    img_ext = base_image.get("ext", "png")
                    if not img_data or len(img_data) < 200:
                        continue

                    # Skip tiny decorative images
                    width = base_image.get("width", 0)
                    height = base_image.get("height", 0)
                    if width < 20 or height < 20:
                        continue

                    img_hash = hashlib.md5(img_data).hexdigest()
                    if img_hash in seen_hashes:
                        continue
                    seen_hashes.add(img_hash)

                    tid = template_id or "unknown"
                    s3_filename = f"img_{page_num + 1}_{img_idx}_{uuid.uuid4().hex[:8]}.{img_ext}"
                    s3_key = s3_service.upload_file(
                        file_bytes=img_data,
                        filename=s3_filename,
                        folder=f"templates/{tenant_id}/{tid}/images"
                    )
                    if not s3_key:
                        continue

                    img_url = s3_service.get_file_url(s3_key, expiration=604800)
                    if not img_url:
                        continue

                    image_blocks.append({
                        "id": f"img_{page_num + 1}_{img_idx}",
                        "type": "image",
                        "src": img_url,
                        "s3_key": s3_key,
                        "alt": f"Page {page_num + 1} image {img_idx + 1}",
                        "page": page_num + 1,
                        "style": {"maxWidth": "100%", "display": "block", "margin": "12px auto"},
                        "editable": False,
                    })
                    logger.info(f"[fitz] Extracted image from page {page_num + 1}: {s3_key} ({width}x{height})")

                except Exception as e:
                    logger.debug(f"[fitz] Failed to extract image xref={xref}: {e}")
                    continue

        doc.close()
        return image_blocks

    def _extract_images_pypdf2(self, file_bytes: bytes, s3_service, tenant_id: str, template_id: str = None) -> List[Dict[str, Any]]:
        """Fallback: Extract images from PDF pages using PyPDF2 and upload to S3."""
        from PyPDF2 import PdfReader

        image_blocks = []
        seen_hashes = set()

        try:
            reader = PdfReader(io.BytesIO(file_bytes))
        except Exception as e:
            logger.warning(f"Failed to read PDF for image extraction: {e}")
            return []

        for page_num, page in enumerate(reader.pages, 1):
            page_images = self._get_page_images(reader, page, page_num)
            for img_idx, (img_data, img_ext) in enumerate(page_images):
                if not img_data or len(img_data) < 200:
                    continue

                img_hash = hashlib.md5(img_data).hexdigest()
                if img_hash in seen_hashes:
                    continue
                seen_hashes.add(img_hash)

                tid = template_id or "unknown"
                s3_filename = f"img_{page_num}_{img_idx}_{uuid.uuid4().hex[:8]}.{img_ext}"
                s3_key = s3_service.upload_file(
                    file_bytes=img_data,
                    filename=s3_filename,
                    folder=f"templates/{tenant_id}/{tid}/images"
                )
                if not s3_key:
                    continue

                img_url = s3_service.get_file_url(s3_key, expiration=604800)
                if not img_url:
                    continue

                image_blocks.append({
                    "id": f"img_{page_num}_{img_idx}",
                    "type": "image",
                    "src": img_url,
                    "s3_key": s3_key,
                    "alt": f"Page {page_num} image {img_idx + 1}",
                    "page": page_num,
                    "style": {"maxWidth": "100%", "display": "block", "margin": "12px auto"},
                    "editable": False,
                })
                logger.info(f"Extracted image from page {page_num}: {s3_key}")

        return image_blocks

    @staticmethod
    def _get_page_images(reader, page, page_num: int) -> List[tuple]:
        """Extract all images from a PDF page via XObject resources."""
        results = []
        try:
            resources = page.get('/Resources')
            if not resources:
                return results
            xobject = resources.get('/XObject')
            if not xobject:
                return results
        except Exception:
            return results

        for key in xobject:
            try:
                obj = xobject[key]
                subtype = str(obj.get('/Subtype', ''))
                if subtype != '/Image':
                    continue

                width = int(obj.get('/Width', 0))
                height = int(obj.get('/Height', 0))
                if width < 10 or height < 10:
                    continue

                data = obj.get_data()
                if not data:
                    continue

                color_space = str(obj.get('/ColorSpace', '/DeviceRGB'))
                filt = obj.get('/Filter')
                filt_str = str(filt) if filt else ''

                if 'DCTDecode' in filt_str:
                    results.append((data, 'jpg'))
                    continue
                if 'JPXDecode' in filt_str:
                    results.append((data, 'jp2'))
                    continue

                try:
                    from PIL import Image as PILImage

                    if '/DeviceRGB' in color_space or 'DeviceRGB' in color_space:
                        mode = 'RGB'
                        expected = width * height * 3
                    elif '/DeviceCMYK' in color_space or 'DeviceCMYK' in color_space:
                        mode = 'CMYK'
                        expected = width * height * 4
                    elif '/DeviceGray' in color_space or 'DeviceGray' in color_space:
                        mode = 'L'
                        expected = width * height
                    else:
                        mode = 'RGB'
                        expected = width * height * 3

                    if len(data) >= expected:
                        img = PILImage.frombytes(mode, (width, height), data[:expected])
                        if mode == 'CMYK':
                            img = img.convert('RGB')
                        buf = io.BytesIO()
                        img.save(buf, format='PNG')
                        results.append((buf.getvalue(), 'png'))
                    else:
                        try:
                            img = PILImage.open(io.BytesIO(data))
                            buf = io.BytesIO()
                            img.save(buf, format='PNG')
                            results.append((buf.getvalue(), 'png'))
                        except Exception:
                            pass
                except ImportError:
                    logger.warning("Pillow not installed — cannot convert raw PDF images")
                except Exception as e:
                    logger.debug(f"Failed to convert image {key}: {e}")
            except Exception as e:
                logger.debug(f"Error processing XObject {key}: {e}")
                continue

        # Also try PyPDF2's built-in .images as a fallback
        try:
            for img in page.images:
                if img.data and len(img.data) > 200:
                    ext = img.name.rsplit('.', 1)[-1].lower() if '.' in img.name else 'png'
                    results.append((img.data, ext))
        except Exception:
            pass

        return results

    @staticmethod
    def _merge_image_blocks(text_blocks: List[Dict[str, Any]], image_blocks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Insert image blocks at the beginning of their respective pages."""
        if not image_blocks:
            return text_blocks

        images_by_page = {}
        for img in image_blocks:
            pg = img.get("page", 1)
            images_by_page.setdefault(pg, []).append(img)

        merged = []
        pages_done = set()

        for block in text_blocks:
            pg = block.get("page", 1)
            if pg not in pages_done and pg in images_by_page:
                merged.extend(images_by_page[pg])
                pages_done.add(pg)
            merged.append(block)

        for pg, imgs in images_by_page.items():
            if pg not in pages_done:
                merged.extend(imgs)

        return merged

    # ── DOCX Conversion ───────────────────────────────────────────────

    def convert_docx_to_blocks(self, file_bytes: bytes, s3_service=None, tenant_id=None, template_id=None) -> List[Dict[str, Any]]:
        """Convert DOCX file to structured content blocks with image extraction.
        Headers/footers are extracted and marked with repeat='per_page' so the
        frontend can repeat them on every page.
        """
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, Emu

        raw_blocks = []
        doc = DocxDocument(io.BytesIO(file_bytes))

        # Extract images from DOCX relationships (body + headers + footers)
        if s3_service and tenant_id:
            docx_images = self._extract_docx_images(doc, s3_service, tenant_id, template_id)
        else:
            docx_images = {}

        # ── Extract header/footer content ───────────────────────────
        header_blocks = []
        footer_blocks = []
        try:
            # Process ONLY the first section's header and footer
            if doc.sections:
                section = doc.sections[0]

                # ── Header extraction ──
                if section.header and section.header.paragraphs:
                    for h_idx, para in enumerate(section.header.paragraphs):
                        text = para.text.strip()
                        has_drawing = False
                        for run in para.runs:
                            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                                has_drawing = True
                                break
                        if has_drawing:
                            for rid, img_block in list(docx_images.items()):
                                if img_block.get("source") == "header":
                                    header_blocks.append({
                                        **img_block,
                                        "id": f"hdr_img_{h_idx}",
                                        "repeat": "per_page",
                                        "position": "header",
                                    })
                                    docx_images.pop(rid, None)
                                    break
                        if text:
                            header_blocks.append({
                                "id": f"hdr_{h_idx}",
                                "type": "paragraph",
                                "content": text,
                                "page": 1,
                                "repeat": "per_page",
                                "position": "header",
                                "style": {"fontSize": "10px", "textAlign": "center", "color": "#666"},
                                "editable": True,
                            })
                    # Extract textbox content from header XML
                    hdr_textbox_texts = self._extract_textbox_content(section.header._element)
                    for t_idx, ttext in enumerate(hdr_textbox_texts):
                        if ttext.strip():
                            header_blocks.append({
                                "id": f"hdr_txbox_{t_idx}",
                                "type": "paragraph",
                                "content": ttext.strip(),
                                "page": 1,
                                "repeat": "per_page",
                                "position": "header",
                                "style": {"fontSize": "9px", "textAlign": "center", "color": "#666"},
                                "editable": True,
                            })

                # ── Footer extraction ──
                if section.footer and section.footer.paragraphs:
                    for f_idx, para in enumerate(section.footer.paragraphs):
                        text = para.text.strip()
                        has_drawing = False
                        for run in para.runs:
                            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                                has_drawing = True
                                break
                        if has_drawing:
                            for rid, img_block in list(docx_images.items()):
                                if img_block.get("source") == "footer":
                                    footer_blocks.append({
                                        **img_block,
                                        "id": f"ftr_img_{f_idx}",
                                        "repeat": "per_page",
                                        "position": "footer",
                                    })
                                    docx_images.pop(rid, None)
                                    break
                        if text:
                            footer_blocks.append({
                                "id": f"ftr_{f_idx}",
                                "type": "paragraph",
                                "content": text,
                                "page": 1,
                                "repeat": "per_page",
                                "position": "footer",
                                "style": {"fontSize": "9px", "textAlign": "center", "color": "#999"},
                                "editable": True,
                            })
                    # Extract textbox content from footer XML (address, phone, etc.)
                    ftr_textbox_texts = self._extract_textbox_content(section.footer._element)
                    for t_idx, ttext in enumerate(ftr_textbox_texts):
                        if ttext.strip():
                            footer_blocks.append({
                                "id": f"ftr_txbox_{t_idx}",
                                "type": "paragraph",
                                "content": ttext.strip(),
                                "page": 1,
                                "repeat": "per_page",
                                "position": "footer",
                                "style": {"fontSize": "9px", "textAlign": "center", "color": "#666"},
                                "editable": True,
                            })
        except Exception as e:
            logger.debug(f"Error extracting header/footer blocks: {e}")

        # Also add any remaining header/footer images that weren't matched to paragraphs
        for rid, img_block in list(docx_images.items()):
            src = img_block.get("source", "")
            if src == "header":
                header_blocks.insert(0, {
                    **img_block,
                    "id": f"hdr_img_extra_{uuid.uuid4().hex[:6]}",
                    "repeat": "per_page",
                    "position": "header",
                })
                docx_images.pop(rid, None)
            elif src == "footer":
                footer_blocks.append({
                    **img_block,
                    "id": f"ftr_img_extra_{uuid.uuid4().hex[:6]}",
                    "repeat": "per_page",
                    "position": "footer",
                })
                docx_images.pop(rid, None)

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Check for inline images in this paragraph
            for run in para.runs:
                if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    for rel in para.part.rels.values():
                        if "image" in (rel.reltype or ""):
                            rid = rel.rId
                            if rid in docx_images:
                                raw_blocks.append(docx_images.pop(rid))
                                break

            if not text:
                continue

            # Determine block type from paragraph style
            style_name = (para.style.name or "").lower()
            block_type = "paragraph"
            heading_level = None
            if "heading 1" in style_name or "title" in style_name:
                block_type = "heading"
                heading_level = 1
            elif "heading 2" in style_name or "subtitle" in style_name:
                block_type = "heading"
                heading_level = 2
            elif "heading 3" in style_name:
                block_type = "heading"
                heading_level = 3
            elif "heading" in style_name:
                block_type = "heading"
                heading_level = 2
            elif "list" in style_name or "bullet" in style_name:
                block_type = "list_item"

            # Detect alignment
            alignment = "left"
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                alignment = "center"
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                alignment = "right"
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                alignment = "justify"

            # Extract formatting from runs
            is_bold = any(run.bold for run in para.runs if run.bold)
            is_italic = any(run.italic for run in para.runs if run.italic)
            font_size = None
            font_family = None
            for run in para.runs:
                if run.font.size:
                    font_size = run.font.size.pt
                if run.font.name:
                    font_family = run.font.name
                if font_size and font_family:
                    break

            # Extract paragraph spacing
            pf = para.paragraph_format
            margin_top = 0
            margin_bottom = 0
            line_height = None
            indent_left = 0
            if pf.space_before and isinstance(pf.space_before, (int, Pt, Emu)):
                try:
                    margin_top = int(pf.space_before.pt) if hasattr(pf.space_before, 'pt') else int(pf.space_before / 12700)
                except Exception:
                    pass
            if pf.space_after and isinstance(pf.space_after, (int, Pt, Emu)):
                try:
                    margin_bottom = int(pf.space_after.pt) if hasattr(pf.space_after, 'pt') else int(pf.space_after / 12700)
                except Exception:
                    pass
            if pf.line_spacing:
                try:
                    ls = float(pf.line_spacing)
                    if ls < 5:  # multiplier (e.g., 1.15, 1.5)
                        line_height = ls
                    else:  # absolute value in Pt
                        line_height = 1.4
                except Exception:
                    pass
            if pf.left_indent:
                try:
                    indent_left = int(pf.left_indent.pt) if hasattr(pf.left_indent, 'pt') else 0
                except Exception:
                    pass

            style = {}
            if font_size:
                style["fontSize"] = f"{int(font_size)}px"
            elif block_type == "heading":
                style["fontSize"] = "24px" if heading_level == 1 else "20px" if heading_level == 2 else "17px"
            else:
                style["fontSize"] = "12px"

            if is_bold or block_type == "heading":
                style["fontWeight"] = "bold"
            if is_italic:
                style["fontStyle"] = "italic"
            if alignment != "left":
                style["textAlign"] = alignment
            if font_family:
                style["fontFamily"] = font_family
            if margin_top:
                style["marginTop"] = f"{margin_top}px"
            if margin_bottom:
                style["marginBottom"] = f"{margin_bottom}px"
            if line_height:
                style["lineHeight"] = str(round(line_height, 2))
            if indent_left:
                style["paddingLeft"] = f"{indent_left}px"

            block = {
                "id": f"blk_{idx}",
                "type": block_type,
                "content": text,
                "page": 1,
                "style": style,
                "editable": True,
            }
            if heading_level:
                block["level"] = heading_level
            raw_blocks.append(block)

        # Handle tables
        for t_idx, table in enumerate(doc.tables):
            rows_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows_data.append(cells)

            if rows_data:
                raw_blocks.append({
                    "id": f"tbl_{t_idx}",
                    "type": "table",
                    "content": rows_data,
                    "page": 1,
                    "style": {"fontSize": "11px"},
                    "editable": True,
                })

        # Append any remaining images that weren't matched to paragraphs
        for img_block in docx_images.values():
            raw_blocks.insert(0, img_block)

        # Group consecutive list_item blocks into list blocks with items array
        blocks = []
        list_buffer = []
        for block in raw_blocks:
            if block.get("type") == "list_item":
                list_buffer.append(block.get("content", ""))
            else:
                if list_buffer:
                    blocks.append({
                        "id": f"list_{len(blocks)}",
                        "type": "list",
                        "items": list_buffer,
                        "page": 1,
                        "style": block.get("style", {"fontSize": "12px"}),
                        "editable": True,
                    })
                    list_buffer = []
                blocks.append(block)
        if list_buffer:
            blocks.append({
                "id": f"list_{len(blocks)}",
                "type": "list",
                "items": list_buffer,
                "page": 1,
                "style": {"fontSize": "12px"},
                "editable": True,
            })

        # Assign page numbers (approximate: ~40 lines per page)
        line_count = 0
        current_page = 1
        for block in blocks:
            content = block.get("content", "")
            if isinstance(content, str):
                lines = max(1, len(content) // 80 + 1)
            elif isinstance(content, list):
                lines = len(content) + 1
            else:
                lines = 1
            line_count += lines
            if line_count > 45:
                current_page += 1
                line_count = lines
            block["page"] = current_page

        # Prepend header blocks and append footer blocks
        # These are marked with repeat='per_page' so the frontend repeats them on each page
        all_blocks = header_blocks + blocks + footer_blocks
        return all_blocks


    @staticmethod
    def _extract_textbox_content(element) -> List[str]:
        """Extract text from WordprocessingML textboxes (wps:txbx) within a DOCX header/footer element.
        These textboxes contain address/phone/contact info that python-docx can't access via paragraph.text.
        """
        texts = []
        seen = set()
        # Iterate all descendant elements looking for txbxContent
        for el in element.iter():
            if el.tag.endswith('}txbxContent') or el.tag == 'txbxContent':
                w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                for para in el:
                    if para.tag == f'{{{w_ns}}}p' or para.tag.endswith('}p'):
                        parts = []
                        for run in para:
                            for t in run:
                                if (t.tag == f'{{{w_ns}}}t' or t.tag.endswith('}t')) and t.text:
                                    parts.append(t.text)
                        line = ''.join(parts).strip()
                        if line and line not in seen:
                            seen.add(line)
                            texts.append(line)
        return texts


    @staticmethod
    def _extract_docx_images(doc, s3_service, tenant_id: str, template_id: str = None) -> Dict[str, Dict]:
        """Extract all images from a DOCX document (body + headers + footers) and upload to S3.
        Returns dict mapping relationship ID → image block.
        """
        image_blocks = {}
        seen_hashes = set()

        # Collect relationships from all parts: main body, headers, footers
        all_rels = []
        # Main body
        for rel_id, rel in doc.part.rels.items():
            all_rels.append((f"body_{rel_id}", rel, "body"))

        # Headers and footers — use unique keys to avoid collisions
        try:
            for sec_idx, section in enumerate(doc.sections):
                for part_name, part in [("header", section.header), ("footer", section.footer)]:
                    if part and hasattr(part, '_element') and part._element is not None:
                        try:
                            for rel_id, rel in part.part.rels.items():
                                unique_key = f"{part_name}_{sec_idx}_{rel_id}"
                                all_rels.append((unique_key, rel, part_name))
                        except Exception:
                            pass
        except Exception as e:
            logger.debug(f"Error extracting header/footer rels: {e}")

        for rel_id, rel, source in all_rels:
            if "image" not in (rel.reltype or ""):
                continue
            try:
                img_data = rel.target_part.blob
                if not img_data or len(img_data) < 200:
                    continue

                img_hash = hashlib.md5(img_data).hexdigest()
                if img_hash in seen_hashes:
                    continue
                seen_hashes.add(img_hash)

                content_type = rel.target_part.content_type or ""
                if "jpeg" in content_type or "jpg" in content_type:
                    ext = "jpg"
                elif "png" in content_type:
                    ext = "png"
                elif "gif" in content_type:
                    ext = "gif"
                else:
                    ext = "png"

                tid = template_id or "unknown"
                s3_filename = f"docx_img_{uuid.uuid4().hex[:8]}.{ext}"
                s3_key = s3_service.upload_file(
                    file_bytes=img_data,
                    filename=s3_filename,
                    folder=f"templates/{tenant_id}/{tid}/images"
                )
                if not s3_key:
                    continue

                img_url = s3_service.get_file_url(s3_key, expiration=604800)
                if not img_url:
                    continue

                image_blocks[rel_id] = {
                    "id": f"docx_img_{uuid.uuid4().hex[:6]}",
                    "type": "image",
                    "src": img_url,
                    "s3_key": s3_key,
                    "alt": "Document image",
                    "page": 1,
                    "source": source,
                    "style": {"maxWidth": "100%", "display": "block", "margin": "12px auto"},
                    "editable": False,
                }
                logger.info(f"Extracted DOCX image from {source}: {s3_key}")

            except Exception as e:
                logger.debug(f"Failed to extract DOCX image {rel_id}: {e}")
                continue

        return image_blocks
